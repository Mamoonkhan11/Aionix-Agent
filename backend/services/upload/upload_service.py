"""
File upload service with text extraction capabilities.

This module provides secure file upload functionality with support for
PDF, DOCX, and TXT files, including text extraction and validation.
"""

import logging
import os
import shutil
from pathlib import Path
from typing import Dict, Optional, Tuple
from uuid import uuid4

import aiofiles
from fastapi import UploadFile

from core.config import settings
from models import IngestionLog, IngestionStatus, IngestionType, RawDocument, DocumentSourceType
from schemas.documents import DocumentCreate

logger = logging.getLogger(__name__)


class FileUploadService:
    """
    Service for handling secure file uploads with text extraction.

    Supports PDF, DOCX, and TXT files with automatic text extraction,
    file validation, and secure storage.
    """

    def __init__(self):
        """Initialize upload service with configuration."""
        self.upload_dir = Path(settings.upload_directory)
        self.max_file_size = settings.max_upload_size
        self.allowed_extensions = set(settings.allowed_file_types)

        # Create upload directory if it doesn't exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file.

        Args:
            file: Uploaded file object

        Raises:
            ValueError: If file is invalid
        """
        # Check file size
        if file.size and file.size > self.max_file_size:
            raise ValueError(
                f"File size {file.size} bytes exceeds maximum allowed size "
                f"of {self.max_file_size} bytes"
            )

        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.allowed_extensions:
            raise ValueError(
                f"File type {file_ext} not allowed. "
                f"Allowed types: {', '.join(self.allowed_extensions)}"
            )

        # Check filename for security
        if not file.filename or file.filename.startswith('.'):
            raise ValueError("Invalid filename")

    async def save_file(self, file: UploadFile) -> Path:
        """
        Save uploaded file to disk with secure filename.

        Args:
            file: Uploaded file object

        Returns:
            Path: Path to saved file

        Raises:
            IOError: If file cannot be saved
        """
        # Generate secure filename
        file_ext = Path(file.filename).suffix.lower()
        secure_filename = f"{uuid4()}{file_ext}"
        file_path = self.upload_dir / secure_filename

        try:
            async with aiofiles.open(file_path, 'wb') as buffer:
                # Read file in chunks to handle large files
                chunk_size = 8192
                while True:
                    chunk = await file.read(chunk_size)
                    if not chunk:
                        break
                    await buffer.write(chunk)

            logger.info(f"File saved successfully: {file_path}")
            return file_path

        except Exception as e:
            # Clean up partial file if save failed
            if file_path.exists():
                file_path.unlink()
            raise IOError(f"Failed to save file: {e}")

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            str: Extracted text

        Raises:
            ValueError: If text extraction fails
        """
        try:
            from PyPDF2 import PdfReader

            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text = ""

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                return text.strip()

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")

    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            str: Extracted text

        Raises:
            ValueError: If text extraction fails
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"

            return text.strip()

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    def extract_text_from_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            str: Extracted text

        Raises:
            ValueError: If text extraction fails
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read().strip()

        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Failed to decode text file with supported encodings")

        except Exception as e:
            raise ValueError(f"Failed to read text file: {e}")

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from file based on its extension.

        Args:
            file_path: Path to file

        Returns:
            str: Extracted text

        Raises:
            ValueError: If file type is not supported or extraction fails
        """
        file_ext = file_path.suffix.lower()

        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    async def process_upload(
        self,
        file: UploadFile,
        user_id: str,
        db_session
    ) -> Tuple[RawDocument, IngestionLog]:
        """
        Process file upload: validate, save, extract text, and create document.

        Args:
            file: Uploaded file
            user_id: ID of user uploading the file
            db_session: Database session

        Returns:
            Tuple[RawDocument, IngestionLog]: Created document and ingestion log

        Raises:
            ValueError: If validation or processing fails
        """
        operation_id = str(uuid4())

        # Create ingestion log
        log_entry = IngestionLog(
            operation_id=operation_id,
            ingestion_type=IngestionType.FILE_UPLOAD,
            status=IngestionStatus.PENDING,
            user_id=user_id,
            parameters={
                "filename": file.filename,
                "content_type": file.content_type,
                "file_size": file.size,
            }
        )

        db_session.add(log_entry)
        await db_session.commit()

        try:
            # Validate file
            self.validate_file(file)

            # Save file to disk
            file_path = await self.save_file(file)

            try:
                # Extract text
                extracted_text = self.extract_text(file_path)

                if not extracted_text.strip():
                    raise ValueError("No text content could be extracted from the file")

                # Create document
                doc_data = DocumentCreate(
                    title=Path(file.filename).stem,  # Filename without extension
                    content=extracted_text,
                    source_type=DocumentSourceType.UPLOAD,
                    metadata={
                        "original_filename": file.filename,
                        "content_type": file.content_type,
                        "file_size": file.size,
                        "uploaded_by": user_id,
                        "file_path": str(file_path),
                        "extraction_method": f"extract_text_from_{file_path.suffix[1:]}",
                        "word_count": len(extracted_text.split()),
                    }
                )

                document = RawDocument(
                    title=doc_data.title,
                    content=doc_data.content,
                    source_type=doc_data.source_type,
                    external_id=str(file_path),  # Use file path as external ID
                    metadata=doc_data.metadata,
                )

                db_session.add(document)

                # Update log
                log_entry.mark_completed(IngestionStatus.SUCCESS)
                log_entry.update_metrics(processed=1, successful=1, failed=0)

                await db_session.commit()
                await db_session.refresh(document)

                logger.info(f"File upload processed successfully: {file.filename}")

                return document, log_entry

            except Exception as e:
                # Clean up file if processing failed
                if file_path.exists():
                    file_path.unlink()
                raise e

        except Exception as e:
            logger.error(f"File upload processing failed: {e}")

            # Update log with failure
            log_entry.record_error(str(e))
            log_entry.update_metrics(processed=1, successful=0, failed=1)

            await db_session.commit()

            raise ValueError(f"Failed to process uploaded file: {e}")

    def cleanup_old_files(self, max_age_days: int = 30) -> int:
        """
        Clean up old uploaded files.

        Args:
            max_age_days: Maximum age of files to keep

        Returns:
            int: Number of files cleaned up
        """
        import time

        current_time = time.time()
        max_age_seconds = max_age_days * 24 * 60 * 60

        cleaned_count = 0
        for file_path in self.upload_dir.glob("*"):
            if file_path.is_file():
                file_age = current_time - file_path.stat().st_mtime
                if file_age > max_age_seconds:
                    try:
                        file_path.unlink()
                        cleaned_count += 1
                    except Exception as e:
                        logger.warning(f"Failed to clean up file {file_path}: {e}")

        logger.info(f"Cleaned up {cleaned_count} old uploaded files")
        return cleaned_count
